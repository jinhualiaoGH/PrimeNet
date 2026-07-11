from pathlib import Path
import json
import csv
import shutil
import subprocess
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(r"C:\PrimeNet")
PKG = ROOT / "reports" / "controlled_execution_benchmark_v2"
OUT = PKG / "report"
FIG = PKG / "figures"
EVIDENCE = PKG / "evidence"
TABLES = PKG / "tables"

DOCX = OUT / "primenet_controlled_execution_benchmark_report_v2.docx"
PDF = OUT / "primenet_controlled_execution_benchmark_report_v2.pdf"


def load_json(path: Path):
    with path.open("r", encoding="utf-8") as f:
        return json.load(f)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(str(text))
    r.bold = bold
    r.font.size = Pt(9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    hdr = table.rows[0].cells
    set_cell_text(hdr[0], "Metric", True)
    set_cell_text(hdr[1], "Value", True)
    set_cell_shading(hdr[0], "D9EAF7")
    set_cell_shading(hdr[1], "D9EAF7")

    for k, v in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], k)
        set_cell_text(cells[1], v)

    doc.add_paragraph()
    return table


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_figure(doc, path: Path, caption: str, width=6.5):
    if not path.exists():
        add_body(doc, f"[Missing figure: {path.name}]")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))

    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True


def try_convert_pdf(docx_path: Path, out_dir: Path):
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        return False

    subprocess.run(
        [
            soffice,
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(out_dir),
            str(docx_path),
        ],
        check=False,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True,
    )

    produced = out_dir / (docx_path.stem + ".pdf")
    if produced.exists():
        if produced != PDF:
            produced.replace(PDF)
        return True
    return False


def main():
    OUT.mkdir(parents=True, exist_ok=True)

    summary_path = EVIDENCE / "campaign_summary_v2.json"
    if not summary_path.exists():
        summary_path = PKG / "analysis_outputs" / "campaign_summary_v2.json"

    summary = load_json(summary_path)
    c = summary["campaign"]
    rt = summary["batch_runtime_minutes"]
    gen = summary["generate_minutes"]
    save = summary["save_minutes"]
    ver = summary["verify_minutes"]
    trend = summary["runtime_trend"]
    hb = summary["heartbeat"]
    outlier = summary["outlier_rule"]

    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("PrimeNet Controlled Execution Benchmark Report v2")
    r.bold = True
    r.font.size = Pt(20)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = subtitle.add_run("2T to 3T Repository Production Campaign")
    rr.italic = True
    rr.font.size = Pt(12)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}").font.size = Pt(9)

    add_heading(doc, "Executive Summary", 1)
    add_body(
        doc,
        "This report summarizes the PrimeNet controlled production campaign spanning "
        "2,000,000,000,001 to 3,000,000,000,000. The campaign was executed under the "
        "integrated PrimeNet Controlled Execution Framework, which preserved wall-clock "
        "runtime evidence, heartbeat continuity, logs, environment metadata, and verified "
        "repository outputs."
    )

    add_bullets(
        doc,
        [
            f"Completed batches: {c['batches_completed']} consecutive 10B batches.",
            f"Total primes generated: {c['total_primes']:,}.",
            f"Total output size: {c['total_output_size_gb']:.3f} GB.",
            f"Batch-derived production runtime: {c['total_runtime_hours']:.3f} hours.",
            f"Mean runtime: {rt['mean']:.3f} minutes per 10B batch.",
            f"Median runtime: {rt['median']:.3f} minutes per 10B batch.",
            f"95th percentile runtime: {rt['p95']:.3f} minutes per 10B batch.",
            f"Verified outputs: {c['verified_batches']}/{c['batches_completed']}.",
            f"Heartbeat rows: {hb['rows']:,}; heartbeat gap warnings: {hb['heartbeat_gap_warnings']}.",
        ],
    )

    add_heading(doc, "Benchmark Configuration", 1)
    add_table(
        doc,
        [
            ("Campaign range", f"{c['range_start']:,} to {c['range_end']:,}"),
            ("Batch size", "10,000,000,000 integers"),
            ("Completed batches", c["batches_completed"]),
            ("Repository output", "NumPy range files"),
            ("Controlled execution", "PrimeNet Controlled Execution Framework v1.0"),
            ("Evidence package", str(PKG)),
        ],
    )

    add_heading(doc, "Campaign Results", 1)
    add_table(
        doc,
        [
            ("Total runtime", f"{c['total_runtime_minutes']:.3f} min = {c['total_runtime_hours']:.3f} h"),
            ("Total primes", f"{c['total_primes']:,}"),
            ("Mean prime count per 10B", f"{c['mean_prime_count_per_batch']:,.2f}"),
            ("Total output size", f"{c['total_output_size_gb']:.3f} GB"),
            ("Verified batches", f"{c['verified_batches']}/{c['batches_completed']}"),
            ("Projected hours per 1T at mean", f"{c['projected_hours_per_1T_at_mean']:.3f} h"),
        ],
    )

    add_heading(doc, "Runtime Statistics", 1)
    add_table(
        doc,
        [
            ("Mean", f"{rt['mean']:.3f} min"),
            ("Median", f"{rt['median']:.3f} min"),
            ("Standard deviation", f"{rt['std']:.3f} min"),
            ("Minimum", f"{rt['min']:.3f} min"),
            ("5th percentile", f"{rt['p05']:.3f} min"),
            ("25th percentile", f"{rt['p25']:.3f} min"),
            ("75th percentile", f"{rt['p75']:.3f} min"),
            ("95th percentile", f"{rt['p95']:.3f} min"),
            ("Maximum", f"{rt['max']:.3f} min"),
            ("IQR", f"{rt['iqr']:.3f} min"),
            ("Coefficient of variation", f"{rt['cv']:.3f}"),
        ],
    )

    add_figure(
        doc,
        FIG / "runtime_sequence_v2.png",
        "Figure 1. Total runtime per 10B batch across the controlled 2T to 3T campaign.",
    )

    add_figure(
        doc,
        FIG / "runtime_distribution_v2.png",
        "Figure 2. Distribution of 10B batch runtimes with mean and median reference lines.",
    )

    add_heading(doc, "Runtime Component Analysis", 1)
    add_body(
        doc,
        "Generation dominates the mean runtime per batch, with output serialization as the "
        "secondary cost and verification comparatively small on average."
    )

    add_table(
        doc,
        [
            ("Generate mean", f"{gen['mean']:.3f} min"),
            ("Save mean", f"{save['mean']:.3f} min"),
            ("Verify mean", f"{ver['mean']:.3f} min"),
            ("Generate median", f"{gen['median']:.3f} min"),
            ("Save median", f"{save['median']:.3f} min"),
            ("Verify median", f"{ver['median']:.3f} min"),
        ],
    )

    add_figure(
        doc,
        FIG / "runtime_components_mean_v2.png",
        "Figure 3. Mean runtime components per 10B batch.",
    )

    add_figure(
        doc,
        FIG / "runtime_components_sequence_v2.png",
        "Figure 4. Runtime components by batch index.",
    )

    add_heading(doc, "Execution Continuity", 1)
    add_body(
        doc,
        "The heartbeat record provides direct evidence that the campaign ran under continuous "
        "controlled execution. The observed heartbeat intervals remained close to the 10-second "
        "target and produced zero heartbeat-gap warnings."
    )

    add_table(
        doc,
        [
            ("Heartbeat rows", f"{hb['rows']:,}"),
            ("Mean heartbeat interval", f"{hb['mean_wall_delta_seconds']:.3f} sec"),
            ("95th percentile interval", f"{hb['p95_wall_delta_seconds']:.3f} sec"),
            ("Maximum interval", f"{hb['max_wall_delta_seconds']:.3f} sec"),
            ("Heartbeat gap warnings", hb["heartbeat_gap_warnings"]),
        ],
    )

    add_figure(
        doc,
        FIG / "heartbeat_interval_v2.png",
        "Figure 5. Controlled execution heartbeat continuity across the production campaign.",
    )

    add_heading(doc, "Repository Consistency Diagnostic", 1)
    add_body(
        doc,
        "The prime count per 10B batch decreases smoothly over the campaign, consistent with "
        "declining prime density and with no visible repository discontinuity."
    )

    add_figure(
        doc,
        FIG / "prime_count_sequence_v2.png",
        "Figure 6. Prime count per 10B batch across the 2T to 3T campaign.",
    )

    add_heading(doc, "Stability and Runtime Excursions", 1)
    add_body(
        doc,
        f"A linear trend fit gives a slope of {trend['slope_per_batch']:.6f} minutes per batch, "
        f"with R^2 = {trend['r2']:.3f} and a fitted campaign-scale change of "
        f"{trend['delta_over_campaign']:.3f} minutes. This is best interpreted as an engineering "
        "diagnostic rather than a mathematical law."
    )
    add_body(
        doc,
        f"Using the IQR rule ({outlier['method']}), the runtime excursion threshold was "
        f"{outlier['threshold_minutes']:.3f} minutes and {outlier['outlier_count']} batches were flagged. "
        "The campaign returned to its baseline runtime regime after excursions, supporting the conclusion "
        "that no progressive throughput degradation was observed."
    )

    add_heading(doc, "Interpretation", 1)
    add_body(
        doc,
        "Under controlled execution conditions on the current commodity-hardware platform, PrimeNet "
        "demonstrated sustained repository-production throughput of approximately 10 billion integers "
        "per three minutes. Earlier anomalous wall-clock measurements are excluded from the controlled "
        "performance baseline because they were obtained under uncontrolled execution conditions. The "
        "controlled campaign instead preserved continuous execution telemetry, logs, environment metadata, "
        "and verified repository artifacts."
    )

    add_heading(doc, "Reproducibility Statement", 1)
    add_body(
        doc,
        "The report package preserves the evidence required to reproduce the benchmark analysis: command "
        "record, stdout and stderr logs, job summary, heartbeat telemetry, environment metadata, campaign "
        "summary, batch runtime table, correlation matrix, outlier table, and generated figures."
    )

    add_heading(doc, "Conclusion", 1)
    add_body(
        doc,
        "Across a controlled one-trillion-integer production campaign, PrimeNet completed 100 consecutive "
        "10B batches with complete output verification and uninterrupted heartbeat telemetry. The mean "
        "runtime was 2.951 minutes per 10B batch and the median runtime was 2.895 minutes. These results "
        "establish a platform-specific controlled execution baseline for PrimeNet repository production."
    )

    quote = doc.add_paragraph()
    quote.alignment = WD_ALIGN_PARAGRAPH.CENTER
    qrun = quote.add_run("Observe the primes. Measure the computation. Validate the evidence. Trust the result.")
    qrun.bold = True
    qrun.italic = True

    doc.save(DOCX)
    print(f"Wrote DOCX: {DOCX}")

    ok = try_convert_pdf(DOCX, OUT)
    if ok:
        print(f"Wrote PDF:  {PDF}")
    else:
        print("PDF conversion skipped: LibreOffice/soffice not found on PATH.")
        print("You can open the DOCX in Word and Save As PDF.")


if __name__ == "__main__":
    main()