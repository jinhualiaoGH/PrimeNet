from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from .captions import FIGURE_CAPTIONS, TABLE_CAPTIONS
from .manuscript_source import section_paragraphs

KEYWORDS = "Computational arithmetic; prime numbers; scientific infrastructure; repository architecture; reproducible computing; computational observatory; research software engineering; analytic number theory validation."

def fmt(value, stats):
    try:
        return value.format(**stats)
    except Exception:
        return value

def abstract_text(stats):
    return [
        "Large-scale computational investigations in number theory frequently require researchers to repeatedly construct computational repositories, implement supporting software, verify computational correctness, organize derived products, and manage increasingly complex computational workflows. While considerable attention has been devoted to algorithms and mathematical analysis, comparatively less attention has been given to the engineering infrastructure required to support persistent, reproducible, and extensible computational investigation.",
        "This paper presents PrimeNet, a reference architecture for persistent computational arithmetic. PrimeNet organizes deterministic repository construction, independent verification, canonical observational coordinates, modular observatories, observation sessions, computational products, atlases, metadata management, registry services, and publication support into a unified computational framework designed for long-term scientific investigation.",
        f"The current reference implementation includes a persistent repository covering the interval {stats['repository_interval']}, containing {stats['verified_prime_numbers']} verified prime numbers organized into {stats['repository_segments']} independently verified repository segments. Exact repository products are additionally compared with classical prime-count, mean-gap, and twin-prime reference laws.",
        "The primary contribution of this work is not a new algorithm for prime computation nor a new mathematical theory of prime numbers. Rather, it is the design, implementation, and validation of a persistent computational architecture that enables computational investigations to become reproducible, extensible, and reusable scientific resources.",
    ]

def _shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)

def _add_table(doc, rows, font_size=8):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            cell = table.cell(i, j)
            cell.text = str(value)
            if i == 0:
                _shade(cell, "EAF4FB")
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)
                    run.bold = i == 0
    return table

def _caption(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.keep_together = True
    run = paragraph.add_run(text)
    run.italic = True
    run.font.size = Pt(8.5)
    return paragraph

def _figure_block(doc, image, caption, width):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    paragraph.add_run().add_picture(str(image), width=Inches(width))
    _caption(doc, caption)

def _configure(doc):
    for section in doc.sections:
        section.top_margin = Inches(.72)
        section.bottom_margin = Inches(.72)
        section.left_margin = Inches(.72)
        section.right_margin = Inches(.72)
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(10.5)
    for name in ["Heading 1", "Heading 2"]:
        styles[name].font.name = "Arial"
        styles[name].font.size = Pt(13 if name == "Heading 1" else 11)

def _add_named_table(doc, table_name, tables, font_size):
    caption = doc.add_paragraph(TABLE_CAPTIONS.get(table_name, table_name))
    caption.paragraph_format.keep_with_next = True
    for run in caption.runs:
        run.bold = True
    rows = tables.get(table_name)
    if rows is None:
        print(f"  [SKIP] manuscript table {table_name} (unavailable)")
    else:
        _add_table(doc, rows, font_size)

def build_docx(root, config, stats, tables):
    root = Path(root)
    fig_dir = root / config["paths"]["figures"]
    out_dir = root / config["paths"]["output"]
    out_dir.mkdir(parents=True, exist_ok=True)

    doc = Document()
    _configure(doc)
    width = float(config.get("layout", {}).get("figure_width_inches", 5.75))
    table_font = config.get("layout", {}).get("table_font_size", 8)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(config.get("title", "PrimeNet"))
    run.bold = True
    run.font.size = Pt(18)

    author = doc.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author.add_run(config.get("author", "Jinhua Liao"))

    affiliation = config.get("affiliation")
    if affiliation:
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run(affiliation)

    doc.add_heading("Abstract", level=1)
    for paragraph in abstract_text(stats):
        doc.add_paragraph(paragraph)
    doc.add_paragraph(f"Keywords: {KEYWORDS}")

    doc.add_heading("PrimeNet Reference Implementation", level=1)
    _add_table(doc, tables["table01_reference_implementation"], table_font)
    doc.add_page_break()

    for section in section_paragraphs(stats):
        doc.add_heading(section["title"], level=1)
        for paragraph in section.get("paragraphs", []):
            doc.add_paragraph(fmt(paragraph, stats))

        figure_names = []
        if section.get("figure"):
            figure_names.append(section["figure"])
        figure_names.extend(section.get("figures", []))
        for figure_name in figure_names:
            image = fig_dir / f"{figure_name}.png"
            if image.exists():
                _figure_block(doc, image, FIGURE_CAPTIONS[figure_name], width)
            else:
                print(f"  [SKIP] manuscript figure {figure_name} (unavailable)")

        table_names = []
        if section.get("table"):
            table_names.append(section["table"])
        table_names.extend(section.get("tables", []))
        for table_name in table_names:
            _add_named_table(doc, table_name, tables, table_font)

        doc.add_page_break()

    doc.add_heading("Appendix A. Canonical Tables", level=1)
    for name, rows in tables.items():
        _add_named_table(doc, name, tables, table_font)
        doc.add_paragraph()

    doc.add_page_break()
    doc.add_heading("Appendix B. Theory-Validation Dataset Summary", level=1)
    doc.add_paragraph(
        "The validation figures and Table 8 are generated from the canonical "
        "theory-validation products identified below. These products preserve "
        "partition-level values for all 300 canonical repository partitions."
    )
    summary_rows = [
        ["Field", "Value"],
        ["Dataset status", stats.get("theory_validation_status", "unavailable")],
        ["Dataset version", stats.get("theory_validation_version", "unavailable")],
        ["Partition count", stats.get("theory_partition_count", "unavailable")],
        ["Numeric domain", f"{stats.get('theory_domain_start', 'unavailable')} – {stats.get('theory_domain_end', 'unavailable')}"],
        ["Exact prime count", stats.get("theory_exact_prime_count", "unavailable")],
        ["Exact twin-prime count", stats.get("theory_exact_twin_count", "unavailable")],
        ["Summary source", stats.get("_theory_validation_source", "unavailable")],
        ["Partition source", stats.get("_theory_partition_source", "unavailable")],
    ]
    _add_table(doc, summary_rows, table_font)

    output = out_dir / "PrimeNet_Architecture_Publication_Draft_v2_4.docx"
    doc.save(output)
    return output
