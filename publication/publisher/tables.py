from pathlib import Path
import csv
from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OPTIONAL_TWIN_PERFORMANCE_KEYS = (
    "twin_numeric_domain", "twin_partitions", "twin_total_gaps",
    "twin_total_events", "twin_global_density",
    "twin_end_to_end_runtime_min", "twin_runtime_accounted_percent",
    "steady_partitions", "steady_total_gaps", "steady_mean_runtime_sec",
    "steady_median_runtime_sec", "steady_runtime_cv_percent",
    "steady_p95_runtime_sec", "steady_gaps_per_sec",
)

THEORY_VALIDATION_KEYS = (
    "theory_exact_prime_count", "theory_exact_twin_count",
    "theory_pnt_prediction", "theory_li_prediction",
    "theory_riemann_r_prediction", "theory_mean_gap",
    "theory_log_midpoint", "theory_hl_prediction",
)

def has_twin_performance(stats):
    return all(stats.get(key) not in (None, "", "N/A") for key in OPTIONAL_TWIN_PERFORMANCE_KEYS)

def has_theory_validation(stats):
    return all(stats.get(key) not in (None, "", "N/A") for key in THEORY_VALIDATION_KEYS)

def make_tables(stats):
    return {
        "table01_reference_implementation": [
            ["Characteristic", "Reference Implementation"],
            ["Repository interval", stats["repository_interval"]],
            ["Verified prime numbers", stats["verified_prime_numbers"]],
            ["Largest stored prime", stats["largest_stored_prime"]],
            ["Repository segments", stats["repository_segments"]],
            ["Segment size", stats["segment_size"]],
            ["Repository verification", stats["repository_verification"]],
            ["Repository construction", stats["repository_construction"]],
            ["Prime Space", stats["prime_space"]],
            ["Observatory Framework", stats["observatory_framework"]],
            ["Product Framework", stats["product_framework"]],
            ["Registry Services", stats["registry_services"]],
        ],
        "table02_design_principles": [
            ["Design Principle", "Architectural Realization"],
            ["Observation before explanation", "Repository + Observatory Framework"],
            ["Persistent repository architecture", "Verified persistent repository segments"],
            ["Reproducibility by design", "Deterministic construction + independent verification"],
            ["Canonical coordinate system", "Prime Space and prime index"],
            ["Separation of responsibilities", "Repository / Observatory / Product / Registry layers"],
            ["Extensibility through modular architecture", "New observatories and products can be added incrementally"],
        ],
        "table03_architectural_components": [
            ["Component", "Primary Responsibility"],
            ["Prime Space", "Canonical observational coordinate framework"],
            ["Repository", "Persistent computational assets"],
            ["Observatory", "Reusable computational investigation"],
            ["Observation Session", "Provenance and execution context"],
            ["Product", "Persistent observational output"],
            ["Atlas", "Curated collection of related products"],
            ["Registry", "Discovery and lifecycle management"],
            ["Publisher", "Publication assets and manuscript generation"],
        ],
        "table04_repository_statistics": [
            ["Metric", "Value"],
            ["Repository interval", stats["repository_interval"]],
            ["Verified primes", stats["verified_prime_numbers"]],
            ["Largest stored prime", stats["largest_stored_prime"]],
            ["Repository segments", stats["repository_segments"]],
            ["Segment size", stats["segment_size"]],
            ["Verification result", stats["repository_verification"]],
        ],
        "table05_software_modules": [
            ["Subsystem", "Responsibility"],
            ["Repository construction services", "Deterministic generation of repository segments"],
            ["Repository management services", "Inventory, catalogs, and lifecycle management"],
            ["Verification services", "Independent repository validation"],
            ["Metadata services", "Repository and product description"],
            ["Observatory execution framework", "Run reusable computational investigations"],
            ["Product management", "Preserve and organize persistent outputs"],
            ["Registry services", "Discover observatories and products"],
            ["Publisher services", "Generate publication figures, tables, manuscript drafts, and publication manifests"],
        ],
        "table06_reproducibility_features": [
            ["Feature", "Purpose"],
            ["Deterministic construction", "Repeatable repository generation"],
            ["Independent verification", "Validate repository integrity"],
            ["Structured metadata", "Preserve repository and product context"],
            ["Observation sessions", "Preserve execution provenance"],
            ["Persistent products", "Enable reuse without recomputation"],
            ["Registry services", "Support discoverability and consistency"],
            ["Publication manifest", "Preserve publication build provenance"],
        ],
        "table07_observational_performance": [
            ["Metric", "Accepted Measurement"],
            ["Scientific domain", stats.get("twin_numeric_domain", "N/A")],
            ["Repository partitions", stats.get("twin_partitions", "N/A")],
            ["Total gaps analyzed", stats.get("twin_total_gaps", "N/A")],
            ["Twin-prime events", stats.get("twin_total_events", "N/A")],
            ["Global twin density", stats.get("twin_global_density", "N/A")],
            ["End-to-end runtime", f"{stats.get('twin_end_to_end_runtime_min', 'N/A')} min"],
            ["Runtime accounted for", stats.get("twin_runtime_accounted_percent", "N/A")],
            ["Conservative steady-state partitions", stats.get("steady_partitions", "N/A")],
            ["Steady-state gaps analyzed", stats.get("steady_total_gaps", "N/A")],
            ["Mean partition runtime", f"{stats.get('steady_mean_runtime_sec', 'N/A')} sec"],
            ["Median partition runtime", f"{stats.get('steady_median_runtime_sec', 'N/A')} sec"],
            ["Runtime coefficient of variation", stats.get("steady_runtime_cv_percent", "N/A")],
            ["Runtime P95", f"{stats.get('steady_p95_runtime_sec', 'N/A')} sec"],
            ["Sustained throughput", f"{stats.get('steady_gaps_per_sec', 'N/A')} gaps/sec"],
        ],
        "table08_classical_validation": [
            ["Quantity", "Exact observation", "Classical reference", "Endpoint relative error"],
            [f"Prime count pi({stats.get('theory_endpoint_x', '3,000,000,000,000')})",
             stats.get("theory_exact_prime_count", "N/A"),
             f"x/log(x) = {stats.get('theory_pnt_prediction', 'N/A')}",
             stats.get("theory_pnt_error_percent", "N/A")],
            ["Prime count", stats.get("theory_exact_prime_count", "N/A"),
             f"Li(x) = {stats.get('theory_li_prediction', 'N/A')}",
             stats.get("theory_li_error_percent", "N/A")],
            ["Prime count", stats.get("theory_exact_prime_count", "N/A"),
             f"R(x) = {stats.get('theory_riemann_r_prediction', 'N/A')}",
             stats.get("theory_riemann_r_error_percent", "N/A")],
            ["Mean outgoing gap", stats.get("theory_mean_gap", "N/A"),
             f"log(midpoint) = {stats.get('theory_log_midpoint', 'N/A')}",
             stats.get("theory_mean_gap_error_percent", "N/A")],
            ["Cumulative twin-prime events", stats.get("theory_exact_twin_count", "N/A"),
             f"Hardy–Littlewood = {stats.get('theory_hl_prediction', 'N/A')}",
             stats.get("theory_hl_error_percent", "N/A")],
        ],
    }

def _write_csv(path, rows):
    with open(path, "w", newline="", encoding="utf-8") as stream:
        csv.writer(stream).writerows(rows)

def _write_markdown(path, rows):
    lines = [
        "| " + " | ".join(rows[0]) + " |",
        "| " + " | ".join(["---"] * len(rows[0])) + " |",
    ]
    lines += ["| " + " | ".join(str(value) for value in row) + " |" for row in rows[1:]]
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")

def _shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)

def _write_docx(path, title, rows):
    doc = Document()
    doc.add_heading(title.replace("_", " ").title(), level=1)
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            cell = table.cell(i, j)
            cell.text = str(value)
            if i == 0:
                _shade(cell, "EAF4FB")
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.size = Pt(8)
                    run.bold = i == 0
    doc.save(path)

def build_all(outdir, stats):
    print("Generating tables...")
    outdir = Path(outdir)
    outdir.mkdir(parents=True, exist_ok=True)
    tables = make_tables(stats)

    if not has_twin_performance(stats):
        tables.pop("table07_observational_performance", None)
        print("  [SKIP] table07_observational_performance (accepted twin-prime performance analysis unavailable)")

    if not has_theory_validation(stats):
        tables.pop("table08_classical_validation", None)
        print("  [SKIP] table08_classical_validation (theory-validation dataset unavailable)")

    for index, (name, rows) in enumerate(tables.items(), 1):
        print(f"  [{index}/{len(tables)}] {name}")
        _write_csv(outdir / f"{name}.csv", rows)
        _write_markdown(outdir / f"{name}.md", rows)
        _write_docx(outdir / f"{name}.docx", name, rows)
    return tables
