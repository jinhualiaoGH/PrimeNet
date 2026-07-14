from pathlib import Path
import csv
from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def make_tables(stats):
 return {
 'table01_reference_implementation':[['Characteristic','Reference Implementation'],['Repository interval',stats['repository_interval']],['Verified prime numbers',stats['verified_prime_numbers']],['Largest stored prime',stats['largest_stored_prime']],['Repository segments',stats['repository_segments']],['Segment size',stats['segment_size']],['Repository verification',stats['repository_verification']],['Repository construction',stats['repository_construction']],['Prime Space',stats['prime_space']],['Observatory Framework',stats['observatory_framework']],['Product Framework',stats['product_framework']],['Registry Services',stats['registry_services']]],
 'table02_design_principles':[['Design Principle','Architectural Realization'],['Observation before explanation','Repository + Observatory Framework'],['Persistent repository architecture','Verified persistent repository segments'],['Reproducibility by design','Deterministic construction + independent verification'],['Canonical coordinate system','Prime Space and prime index'],['Separation of responsibilities','Repository / Observatory / Product / Registry layers'],['Extensibility through modular architecture','New observatories and products can be added incrementally']],
 'table03_architectural_components':[['Component','Primary Responsibility'],['Prime Space','Canonical observational coordinate framework'],['Repository','Persistent computational assets'],['Observatory','Reusable computational investigation'],['Observation Session','Provenance and execution context'],['Product','Persistent observational output'],['Atlas','Curated collection of related products'],['Registry','Discovery and lifecycle management'],['Publisher','Publication assets and manuscript generation']],
 'table04_repository_statistics':[['Metric','Value'],['Repository interval',stats['repository_interval']],['Verified primes',stats['verified_prime_numbers']],['Largest stored prime',stats['largest_stored_prime']],['Repository segments',stats['repository_segments']],['Segment size',stats['segment_size']],['Verification result',stats['repository_verification']]],
 'table05_software_modules':[['Subsystem','Responsibility'],['Repository construction services','Deterministic generation of repository segments'],['Repository management services','Inventory, catalogs, and lifecycle management'],['Verification services','Independent repository validation'],['Metadata services','Repository and product description'],['Observatory execution framework','Run reusable computational investigations'],['Product management','Preserve and organize persistent outputs'],['Registry services','Discover observatories and products'],['Publisher services','Generate publication figures, tables, manuscript drafts, and publication manifests']],
 'table06_reproducibility_features':[['Feature','Purpose'],['Deterministic construction','Repeatable repository generation'],['Independent verification','Validate repository integrity'],['Structured metadata','Preserve repository and product context'],['Observation sessions','Preserve execution provenance'],['Persistent products','Enable reuse without recomputation'],['Registry services','Support discoverability and consistency'],['Publication manifest','Preserve publication build provenance']],
 'table07_observational_performance':[
     ['Metric','Accepted Measurement'],
     ['Scientific domain',stats['twin_numeric_domain']],
     ['Repository partitions',stats['twin_partitions']],
     ['Total gaps analyzed',stats['twin_total_gaps']],
     ['Twin-prime events',stats['twin_total_events']],
     ['Global twin density',stats['twin_global_density']],
     ['End-to-end runtime',f"{stats['twin_end_to_end_runtime_min']} min"],
     ['Runtime accounted for',stats['twin_runtime_accounted_percent']],
     ['Conservative steady-state partitions',stats['steady_partitions']],
     ['Steady-state gaps analyzed',stats['steady_total_gaps']],
     ['Mean partition runtime',f"{stats['steady_mean_runtime_sec']} sec"],
     ['Median partition runtime',f"{stats['steady_median_runtime_sec']} sec"],
     ['Runtime coefficient of variation',stats['steady_runtime_cv_percent']],
     ['Runtime P95',f"{stats['steady_p95_runtime_sec']} sec"],
     ['Sustained throughput',f"{stats['steady_gaps_per_sec']} gaps/sec"],
 ],}
def _write_csv(path,rows):
 with open(path,'w',newline='',encoding='utf-8') as f: csv.writer(f).writerows(rows)
def _write_markdown(path,rows):
 lines=['| '+' | '.join(rows[0])+' |','| '+' | '.join(['---']*len(rows[0]))+' |']; lines += ['| '+' | '.join(r)+' |' for r in rows[1:]]; path.write_text('\n'.join(lines)+'\n',encoding='utf-8')
def _shade(cell,fill):
 tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),fill); tcPr.append(shd)
def _write_docx(path,title,rows):
 doc=Document(); doc.add_heading(title.replace('_',' ').title(),level=1); table=doc.add_table(rows=len(rows),cols=len(rows[0])); table.alignment=WD_TABLE_ALIGNMENT.CENTER; table.style='Table Grid'
 for i,row in enumerate(rows):
  for j,val in enumerate(row):
   cell=table.cell(i,j); cell.text=str(val)
   if i==0: _shade(cell,'EAF4FB')
   for p in cell.paragraphs:
    p.alignment=WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs: run.font.size=Pt(8); run.bold = i==0
 doc.save(path)
def build_all(outdir,stats):
 print('Generating tables...'); outdir=Path(outdir); outdir.mkdir(parents=True,exist_ok=True); tables=make_tables(stats)
 for i,(name,rows) in enumerate(tables.items(),1): print(f'  [{i}/{len(tables)}] {name}'); _write_csv(outdir/f'{name}.csv',rows); _write_markdown(outdir/f'{name}.md',rows); _write_docx(outdir/f'{name}.docx',name,rows)
 return tables
