from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from .captions import FIGURE_CAPTIONS,TABLE_CAPTIONS
from .manuscript_source import section_paragraphs
KEYWORDS='Computational arithmetic; prime numbers; scientific infrastructure; repository architecture; reproducible computing; computational observatory; research software engineering.'
def fmt(s,stats):
 try: return s.format(**stats)
 except Exception: return s
def abstract_text(stats):
 return [
  'Large-scale computational investigations in number theory frequently require researchers to repeatedly construct computational repositories, implement supporting software, verify computational correctness, organize derived products, and manage increasingly complex computational workflows. While considerable attention has been devoted to algorithms and mathematical analysis, comparatively less attention has been given to the engineering infrastructure required to support persistent, reproducible, and extensible computational investigation.',
  'This paper presents PrimeNet, a reference architecture for persistent computational arithmetic. PrimeNet organizes deterministic repository construction, independent verification, canonical observational coordinates, modular observatories, observation sessions, computational products, atlases, metadata management, registry services, and publication support into a unified computational framework designed for long-term scientific investigation.',
  f"The current reference implementation includes a persistent repository covering the interval {stats['repository_interval']}, containing {stats['verified_prime_numbers']} verified prime numbers organized into {stats['repository_segments']} independently verified repository segments, with complete repository verification performed after deterministic construction.",
  'The primary contribution of this work is not a new algorithm for prime computation nor a new mathematical theory of prime numbers. Rather, it is the design, implementation, and validation of a persistent computational architecture that enables computational investigations to become reproducible, extensible, and reusable scientific resources.' ]
def _shade(cell,fill):
 tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),fill); tcPr.append(shd)
def _add_table(doc,rows,font_size=8):
 table=doc.add_table(rows=len(rows),cols=len(rows[0])); table.style='Table Grid'
 for i,row in enumerate(rows):
  for j,val in enumerate(row):
   cell=table.cell(i,j); cell.text=str(val)
   if i==0: _shade(cell,'EAF4FB')
   for p in cell.paragraphs:
    p.alignment=WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs: run.font.size=Pt(font_size); run.bold = i==0
 return table
def _caption(doc,text):
 p=doc.add_paragraph(); p.paragraph_format.keep_together=True; r=p.add_run(text); r.italic=True; r.font.size=Pt(8.5); return p
def _figure_block(doc,img,caption,width):
 p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.keep_with_next=True; p.add_run().add_picture(str(img),width=Inches(width)); _caption(doc,caption)
def _configure(doc):
 for section in doc.sections:
  section.top_margin=Inches(.72); section.bottom_margin=Inches(.72); section.left_margin=Inches(.72); section.right_margin=Inches(.72)
 styles=doc.styles; styles['Normal'].font.name='Times New Roman'; styles['Normal'].font.size=Pt(10.5)
 for s in ['Heading 1','Heading 2']:
  styles[s].font.name='Arial'; styles[s].font.size=Pt(13 if s=='Heading 1' else 11)
def build_docx(root,config,stats,tables):
 root=Path(root); fig_dir=root/config['paths']['figures']; out_dir=root/config['paths']['output']; out_dir.mkdir(parents=True,exist_ok=True)
 doc=Document(); _configure(doc); width=float(config.get('layout',{}).get('figure_width_inches',5.75)); table_font=config.get('layout',{}).get('table_font_size',8)
 title=doc.add_paragraph(); title.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=title.add_run(config.get('title','PrimeNet')); r.bold=True; r.font.size=Pt(18)
 author=doc.add_paragraph(); author.alignment=WD_ALIGN_PARAGRAPH.CENTER; author.add_run(config.get('author','Jinhua Liao'))
 aff=config.get('affiliation')
 if aff:
  p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run(aff)
 doc.add_heading('Abstract',level=1)
 for para in abstract_text(stats): doc.add_paragraph(para)
 doc.add_paragraph(f'Keywords: {KEYWORDS}')
 doc.add_heading('PrimeNet Reference Implementation',level=1); _add_table(doc,tables['table01_reference_implementation'],table_font)
 doc.add_page_break()
 for sec in section_paragraphs(stats):
  doc.add_heading(sec['title'],level=1)
  for para in sec.get('paragraphs',[]): doc.add_paragraph(fmt(para,stats))
  if sec.get('figure'):
   img=fig_dir/f"{sec['figure']}.png"
   if img.exists(): _figure_block(doc,img,FIGURE_CAPTIONS[sec['figure']],width)
  if sec.get('table'):
   cap=doc.add_paragraph(TABLE_CAPTIONS.get(sec['table'],sec['table'])); cap.paragraph_format.keep_with_next=True
   for rr in cap.runs: rr.bold=True
   _add_table(doc,tables[sec['table']],table_font)
  doc.add_page_break()
 doc.add_heading('Appendix A. Canonical Tables',level=1)
 for name,rows in tables.items():
  cap=doc.add_paragraph(TABLE_CAPTIONS.get(name,name)); cap.paragraph_format.keep_with_next=True
  for rr in cap.runs: rr.bold=True
  _add_table(doc,rows,table_font); doc.add_paragraph()
 out=out_dir/'PrimeNet_Architecture_Publication_Draft_v2_2.docx'; doc.save(out); return out
